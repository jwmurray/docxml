"""python-docx benchmark, for comparison against `cargo run --example bench --release`.

Times, per fixture in tests/fixtures/:
  - Document(path) open, N times
  - open + save (to an in-memory BytesIO) round trip, N times
Plus interpreter + `import docx` startup time, measured separately by
spawning subprocesses (docxml's equivalent operations have no per-process
startup cost worth measuring, since it's a compiled binary; python-docx's
does, and it dominates short-lived scripts, so it is reported on its own).

python-docx is slower per operation than the compiled Rust code, so N is
smaller here (50 vs 200) to keep total run time reasonable.

Run with:
    uv run --with python-docx python benches/bench_python_docx.py
"""

import io
import statistics
import subprocess
import sys
import time
from pathlib import Path

import docx

FIXTURES_DIR = Path(__file__).resolve().parent.parent / "tests" / "fixtures"
N_DOC_OPS = 50
N_STARTUP = 10


def median_ms(durations: list[float]) -> float:
    return statistics.median(durations) * 1000.0


def bench_open(path: Path, n: int = N_DOC_OPS) -> float:
    durations = []
    for _ in range(n):
        start = time.perf_counter()
        docx.Document(str(path))
        durations.append(time.perf_counter() - start)
    return median_ms(durations)


def bench_open_save(path: Path, n: int = N_DOC_OPS) -> float:
    durations = []
    for _ in range(n):
        start = time.perf_counter()
        document = docx.Document(str(path))
        buf = io.BytesIO()
        document.save(buf)
        durations.append(time.perf_counter() - start)
    return median_ms(durations)


def bench_startup(n: int = N_STARTUP) -> float:
    """Median wall time of a fresh `python -c "import docx"` subprocess.

    Uses the current interpreter (the one `uv run --with python-docx`
    prepared), so the import actually resolves python-docx.
    """
    durations = []
    for _ in range(n):
        start = time.perf_counter()
        subprocess.run(
            [sys.executable, "-c", "import docx"],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        durations.append(time.perf_counter() - start)
    return median_ms(durations)


def main() -> None:
    fixtures = sorted(FIXTURES_DIR.glob("*.docx"))
    print(
        f"python-docx bench — {len(fixtures)} fixture(s), "
        f"N={N_DOC_OPS} per doc operation, reporting median milliseconds/op\n"
    )
    print(f"{'fixture':<22} {'open (ms)':>12} {'open+save (ms)':>16}")
    print("-" * 52)
    for fixture in fixtures:
        open_ms = bench_open(fixture)
        roundtrip_ms = bench_open_save(fixture)
        print(f"{fixture.name:<22} {open_ms:>12.3f} {roundtrip_ms:>16.3f}")

    print()
    startup_ms = bench_startup()
    print(
        f"interpreter + `import docx` startup: {startup_ms:.2f} ms "
        f"(median of {N_STARTUP} subprocess runs)"
    )


if __name__ == "__main__":
    main()
