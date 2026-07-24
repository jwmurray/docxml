"""Generate tests/fixtures/styles_toc.docx.

Multiple heading levels, a custom character style, and a PAGE field code in
the footer (raw OXML injection, since python-docx has no field API).

Run with:
    uv run --with python-docx python tests/fixtures/generators/styles_toc.py
"""

from pathlib import Path

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIXTURES_DIR = Path(__file__).resolve().parent.parent
OUTPUT_PATH = FIXTURES_DIR / "styles_toc.docx"


def add_field(paragraph, instruction: str) -> None:
    """Insert a simple Word field (e.g. `PAGE`) via raw OXML.

    Built from the three-run sequence Word itself emits: a `begin` fldChar, a
    run carrying the field instruction text, and an `end` fldChar. Skipped
    gracefully by callers if it ever proves fragile, per the task brief, but
    this minimal shape has been stable across Word versions for `PAGE`/`NUMPAGES`.
    """
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run3._r.append(separate)

    # Cached display value shown before the field is updated by Word.
    run4 = paragraph.add_run("1")

    run5 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run5._r.append(end)


def build() -> None:
    document = docx.Document()

    # Custom character style.
    styles = document.styles
    emphasis_style = styles.add_style("FixtureEmphasis", WD_STYLE_TYPE.CHARACTER)
    emphasis_style.font.bold = True
    emphasis_style.font.italic = True
    emphasis_style.font.color.rgb = docx.shared.RGBColor(0xB0, 0x00, 0x20)

    document.add_heading("Styles and Fields Fixture", level=0)

    document.add_heading("Chapter One", level=1)
    p = document.add_paragraph("This chapter introduces the ")
    run = p.add_run("custom character style")
    run.style = emphasis_style
    p.add_run(" used throughout this fixture.")

    document.add_heading("Section 1.1", level=2)
    document.add_paragraph("Body text under a level-2 heading.")

    document.add_heading("Section 1.1.1", level=3)
    document.add_paragraph("Body text under a level-3 heading.")

    document.add_heading("Chapter Two", level=1)
    document.add_paragraph("A second top-level chapter, for heading-level variety.")

    document.add_heading("Section 2.1", level=2)
    document.add_paragraph("More body text.")

    # Footer with a PAGE field code.
    section = document.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Page "
    add_field(footer_paragraph, "PAGE")

    document.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
