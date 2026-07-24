"""Generate tests/fixtures/hyperlinks_images.docx.

A document with an external hyperlink and an inline image. The image is a
tiny 4x4 red PNG built by hand (signature + IHDR/IDAT/IEND chunks) so the
generator needs nothing beyond python-docx.

Run with:
    uv run --with python-docx python tests/fixtures/generators/hyperlinks_images.py
"""

import struct
import zlib
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches

FIXTURES_DIR = Path(__file__).resolve().parent.parent
OUTPUT_PATH = FIXTURES_DIR / "hyperlinks_images.docx"


def make_png(width: int, height: int, rgb: tuple[int, int, int]) -> bytes:
    """Build a minimal truecolor PNG of a solid color, no external deps."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    signature = b"\x89PNG\r\n\x1a\n"
    # 8-bit depth, color type 2 (truecolor), no filter/interlace.
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw_row = bytes([0]) + bytes(rgb) * width  # filter-type byte + pixels
    raw = raw_row * height
    idat = zlib.compress(raw)
    return signature + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def add_hyperlink(paragraph, url: str, text: str):
    """Add an external hyperlink run to `paragraph` via raw OXML injection.

    python-docx has no public write API for hyperlinks (only read, via
    `Paragraph.hyperlinks`), so this follows the well-known recipe: register
    an external relationship on the part, then build the `w:hyperlink`
    element by hand.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})

    run = paragraph._p.makeelement(qn("w:r"), {})
    rpr = run.makeelement(qn("w:rPr"), {})
    rstyle = rpr.makeelement(qn("w:rStyle"), {qn("w:val"): "Hyperlink"})
    rpr.append(rstyle)
    run.append(rpr)

    t = run.makeelement(qn("w:t"), {})
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def build() -> None:
    document = docx.Document()

    document.add_heading("Hyperlinks and Images Fixture", level=1)

    p = document.add_paragraph("See the docxml project on ")
    add_hyperlink(p, "https://github.com/jwmurray/docxml", "GitHub")
    p.add_run(" for details.")

    document.add_paragraph("An inline image follows:")

    png_bytes = make_png(4, 4, (220, 20, 20))
    png_path = FIXTURES_DIR / "_tmp_red.png"
    png_path.write_bytes(png_bytes)
    try:
        document.add_picture(str(png_path), width=Inches(1))
    finally:
        png_path.unlink(missing_ok=True)

    document.add_paragraph("End of fixture.")

    document.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
