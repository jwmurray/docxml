"""Generate tests/fixtures/tables_merged.docx.

Tables with merged cells (horizontal and vertical), a nested table inside a
cell, and varied column widths.

Run with:
    uv run --with python-docx python tests/fixtures/generators/tables_merged.py
"""

from pathlib import Path

import docx
from docx.shared import Inches

FIXTURES_DIR = Path(__file__).resolve().parent.parent
OUTPUT_PATH = FIXTURES_DIR / "tables_merged.docx"


def build() -> None:
    document = docx.Document()

    document.add_heading("Merged and Nested Tables Fixture", level=1)

    # --- Table 1: horizontal + vertical merges, varied column widths -------
    document.add_paragraph("Table with horizontal and vertical merges:")
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    widths = [Inches(0.75), Inches(2.0), Inches(1.0), Inches(1.25)]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    # Header row.
    header_cells = table.rows[0].cells
    for i, text in enumerate(["ID", "Description", "Qty", "Price"]):
        header_cells[i].text = text

    # Row 1: merge "Description" and "Qty" cells horizontally.
    table.cell(1, 0).text = "1"
    merged = table.cell(1, 1).merge(table.cell(1, 2))
    merged.text = "Widget (horizontal merge over Description+Qty)"
    table.cell(1, 3).text = "$10.00"

    # Rows 2-3: merge the "ID" column vertically.
    table.cell(2, 0).text = "2"
    table.cell(3, 0).text = ""
    vmerged = table.cell(2, 0).merge(table.cell(3, 0))
    vmerged.text = "2 (vertical merge)"

    table.cell(2, 1).text = "Gadget"
    table.cell(2, 2).text = "3"
    table.cell(2, 3).text = "$5.00"
    table.cell(3, 1).text = "Gadget, part two"
    table.cell(3, 2).text = "1"
    table.cell(3, 3).text = "$7.50"

    document.add_paragraph()

    # --- Table 2: a table nested inside a cell of an outer table -----------
    document.add_paragraph("Table with a nested table inside a cell:")
    outer = document.add_table(rows=2, cols=2)
    outer.style = "Table Grid"
    outer.cell(0, 0).text = "Outer cell A"
    outer.cell(0, 1).text = "Outer cell B (contains nested table below)"

    nested_host = outer.cell(1, 1)
    # Clear the default empty paragraph's placeholder text before nesting.
    nested_host.paragraphs[0].text = "Nested table:"
    nested = nested_host.add_table(rows=2, cols=2)
    nested.style = "Table Grid"
    nested.cell(0, 0).text = "n00"
    nested.cell(0, 1).text = "n01"
    nested.cell(1, 0).text = "n10"
    nested.cell(1, 1).text = "n11"

    outer.cell(1, 0).text = "Outer cell C"

    document.add_paragraph("End of fixture.")

    document.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
